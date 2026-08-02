#!/usr/bin/env python3
"""Generate the Amorphous Applications / Hermes Station overview docx.

Regenerated to match the shipped system (React SPA + real AIAgent + live-update
engine + 13 components), not the v1 PoC. Source of truth: DEVELOPER_GUIDE.md.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from pathlib import Path

doc = Document()

# base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)


def p(text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    return para


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(headers, rows, style_name="Light Grid Accent 1"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = style_name
    hc = t.rows[0].cells
    for i, htext in enumerate(headers):
        hc[i].text = htext
        for para in hc[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for row in rows:
        rc = t.add_row().cells
        for i, cell in enumerate(row):
            rc[i].text = cell
    return t


# ---------------------------------------------------------------- title
doc.add_heading("Hermes Station", 0)
p("Amorphous Applications — a self-evolving, Hermes-powered mission-control dashboard", bold=True)
p("Nous Research · Proof of Concept · demos/amorphous in the hermes-agent repo "
  "(standalone: NousResearch/hermes-station)")

# ---------------------------------------------------------------- 1
doc.add_heading("1. The Idea", 1)
p("Every person, team, or company gets a personal mission-control surface that is "
  "HERMES-powered and HERMES-shaped: the agent builds it, watches how it is used, and "
  "continuously reshapes it around the user's actual needs. It surfaces agent activity, "
  "one-click repeatable workflows, and internal + external datasources (public and "
  "private), with the goal of becoming the only interaction layer the user needs to do "
  "their job. It is 'amorphous' because it has no fixed final form — it evolves with usage.")
p("Two properties define the build: the chat dock is a REAL Hermes AIAgent (the user's "
  "configured model + toolsets) that reads and reshapes the board live, and every "
  "datasource shows REAL data — an unconnected source renders setup instructions, never "
  "fake numbers.")

# ---------------------------------------------------------------- 2
doc.add_heading("2. What the PoC Demonstrates", 1)
bullets([
    "A local web server (FastAPI) serving a per-user React single-page app composed from a "
    "typed component library; the built SPA is served directly by the backend.",
    "Onboarding: a template gallery (Developer / Trader / Executive / Blank), a live scan of "
    "the machine's real connections (gh CLI auth, local git repos, market/RSS/weather, plus "
    "key-gated Datadog/Better Stack), and a free-form brief the agent uses to tailor the "
    "board before first render.",
    "An invariant agent chat dock — a full AIAgent, docked bottom or right, or detached into "
    "a draggable/resizable floating window, but always present.",
    "Per-component scoped chat: right-click any card to talk to an agent that can only see "
    "and mutate that one component (blast radius enforced server-side).",
    "A drag-and-resize grid: cards move and resize freely; user positions are saved verbatim "
    "and never re-packed by the server.",
    "Pop-out: every card maximizes into a dialog with a full-size view and an 'Ask Hermes' tab.",
    "A command palette (Cmd/Ctrl-K): fuzzy jump-to-card, run-workflow, and station actions, "
    "with free text falling through to the main agent.",
    "Workflow components with a full typed-input form system (text, textarea, number, select, "
    "switch, slider, date, password) that trigger repeatable Hermes tasks and render results "
    "as native rich text.",
    "A live-update engine: a server-side watcher re-queries each visible component at its own "
    "cadence and pushes changes over SSE, so cards update the moment reality changes.",
    "Full-fidelity interaction telemetry feeding an evolution curator that proposes dashboard "
    "mutations on a schedule — never auto-applied — through a preview + approval + feedback tray.",
    "Chat-prompted rebuild: '/rebuild <describe it>' makes the agent redesign the whole board "
    "as a reviewable proposal.",
    "Versioned layout history: every applied change is a new version tagged with its source "
    "(seed / user / agent / curator / rebuild), so evolution is auditable and reversible.",
])

# ---------------------------------------------------------------- 3
doc.add_heading("3. Architecture & Primitives", 1)
p("Backend is Python (FastAPI) so it can embed a real Hermes AIAgent in-process; the "
  "frontend is a TypeScript React SPA. They meet over a REST + SSE contract.")
table(
    ["Primitive", "Responsibility"],
    [
        ("Layout store (store.py)",
         "SQLite. Append-only versioned dashboard specs per user with source attribution; "
         "interaction telemetry; curator proposals; workflows + run history; user feedback."),
        ("Component library (components.py)",
         "13 component types + the mutation engine (promote/shrink/resize/hide/show/remove/"
         "retitle/add/set_props/set_notes/move_chat_dock/replace_spec) + position-preserving "
         "placement + per-template seeds."),
        ("Datasources (datasources.py)",
         "Uniform query(source, props) -> typed payload. REAL data only: git, GitHub (gh CLI), "
         "system stats, crypto (CoinGecko), RSS, weather (Open-Meteo), Datadog, Better Stack, "
         "commit heatmap, live log tail, station activity. Unconnected -> setup instructions. "
         "Per-source refresh cadence defaults."),
        ("Agent bridge (agent_bridge.py)",
         "Builds a real AIAgent from the user's resolved provider/model with terminal, file, "
         "web, vision, skills + five station_* tools. Two session scopes (whole-board vs "
         "single-component), enforced server-side."),
        ("Evolution curator (curator.py)",
         "Heuristic engine (always) + LLM refinement (when live) + rejection memory: scores "
         "component usage, promotes hot panels, shrinks/hides cold ones, mints workflow "
         "shortcuts from repeated chat prompts, and never re-proposes recently rejected changes."),
        ("Live data watcher (server.py)",
         "Re-queries every visible component at its cadence, content-hashes results, and pushes "
         "a targeted 'data_changed' SSE event so exactly the changed card refetches."),
        ("Server + SPA (server.py, web/)",
         "REST API + SSE, curator scheduler, client black box, and the built React app "
         "(Vite + Tailwind v4 + Radix + recharts + react-grid-layout)."),
    ],
)

p("Station tools the embedded agent can call:", bold=True)
bullets([
    "station_get_dashboard — read the layout (or just one component, when scoped).",
    "station_mutate — apply layout mutations; chat-driven edits apply immediately (the user "
    "is watching), only curator/rebuild proposals go through the approval tray.",
    "station_query_datasource — dry-run any source so the agent verifies data before wiring a card.",
    "station_create_workflow — mint a reusable workflow, including a typed input form.",
    "station_component_data — read exactly what the user currently sees in a card.",
])

# ---------------------------------------------------------------- 4
doc.add_heading("4. The Component Library (13 types)", 1)
table(
    ["Type", "What it shows"],
    [
        ("metric", "A single figure with a direction-aware delta chip; type scales to the card."),
        ("timeseries", "recharts area chart with a soft neon glow and gradient fill."),
        ("table", "Sortable columns, status badges, author avatars; renders the rows that fit."),
        ("kv", "Key/value pairs; threshold-colored usage bars; multi-column when wide."),
        ("feed", "An icon-rail activity timeline."),
        ("links", "Favicon link lists; multi-column when wide."),
        ("workflow_button / workflow_panel",
         "Run a Hermes workflow; typed input form; result rendered as rich text."),
        ("notes", "Markdown notes (the agent-maintained briefing) as native prose."),
        ("connections", "Live connection status with enabled/off state."),
        ("heatmap", "GitHub-style commit calendar with month labels, weekday rail, and legend."),
        ("logs", "Live file tail, severity-colored, follows only when scrolled to the bottom."),
        ("tasklist", "An agent-editable checklist — Hermes tracks its own work on the board."),
    ],
)

# ---------------------------------------------------------------- 5
doc.add_heading("5. The Feedback Loops", 1)
p("Four loops make the dashboard amorphous:")
bullets([
    "Instant — ask the main chat, the agent calls station_mutate, the board updates live over SSE.",
    "Scoped — per-card chat, with a server-enforced blast radius of one component.",
    "Slow — telemetry accumulates; the curator (heuristics + LLM + rejection memory) proposes "
    "mutations; the user previews (Try it), then approves or rejects with feedback that steers "
    "the next run; every change is a new, reversible layout version.",
    "Data — the watcher hashes source payloads and pushes targeted refreshes the moment data changes.",
])

# ---------------------------------------------------------------- 6
doc.add_heading("6. Design", 1)
p("A named design system applied literally, derived from the reference mockup: deep navy "
  "canvas, slate cards with hairline borders, a single electric-blue accent (emerald/amber/"
  "red reserved for data semantics), Inter with uppercase micro-labels and tabular numerals. "
  "Chrome is structural — a left sidebar, a stats strip, a scrollable grid, and an in-flow "
  "chat console — never floating boxes. Controls are built on Radix primitives (shadcn "
  "architecture); no browser-default widget chrome leaks through. Agent text renders as rich "
  "prose, never raw markdown. A classical blue-ink engraved Hermes identity (bust + winged "
  "helm) runs through the sidebar, inspector, and empty states, and a capability radar chart "
  "anchors a right-hand inspector rail.")

# ---------------------------------------------------------------- 7
doc.add_heading("7. Roadmap: PoC -> Product", 1)
bullets([
    "Webhook ingest (GitHub, Datadog) feeding the same data_changed event shape — push, not poll.",
    "Streaming agent turns into the dock (token deltas rather than request/response).",
    "A design-token constraint on station_mutate so agent-driven evolution can never emit "
    "off-system styling.",
    "Bespoke agent-generated components in a sandboxed iframe contract — the fully amorphous tier.",
    "Multi-user / org: shared boards with role-scoped sections, hosted on Nous Portal.",
    "Package the agent bridge so the standalone repo no longer depends on a sibling hermes-agent checkout.",
])

# ---------------------------------------------------------------- 8
doc.add_heading("8. Running It", 1)
p("cd demos/amorphous/web && npm install && npm run build")
p("python demos/amorphous/serve.py --port 8877 --db ~/.hermes/hermes-station.db --curator-minutes 60")
p("Open http://localhost:8877 — first visit is onboarding. simulate.py synthesizes a period "
  "of usage for curator demos; web/record-demo.mjs records an end-to-end demo video.")

out = Path(__file__).parent / "Amorphous-Applications-PoC.docx"
doc.save(out)
print(f"wrote {out}")
