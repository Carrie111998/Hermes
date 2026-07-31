#!/usr/bin/env python3
"""Embed one comic illustration below each formula quote in the booklet docx, and patch Markdown with image links."""

import sys, os, re
sys.path.insert(0, r'C:\Users\juzhiyuan\AppData\Local\Programs\Python\Python312\Lib\site-packages')
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = r'D:\Users\juzhiyuan\Pictures\Screenshots\盲派十干直断口诀小册子'
IN_DOC = os.path.join(BASE, '盲派十干直断口诀_五种实战用法_v2.docx')
OUT_DOC = os.path.join(BASE, '盲派十干直断口诀_五种实战用法_v3_逐句漫画版.docx')
COMIC_DIR = os.path.join(BASE, 'illustrations', 'comic')
MD_PATH = r'F:\BaiduSyncdisk\炁运堂盲派命理\《盲派十干直断口诀》五种实战用法.md'
OUT_MD = r'F:\BaiduSyncdisk\炁运堂盲派命理\《盲派十干直断口诀》五种实战用法_逐句漫画版.md'

# Formula text in docx paragraphs often appears with line breaks. Use compact text matching.
FORMULA_TO_IMAGE = {
    '四柱排定看日干日干断诀记心间': '总论1.png',
    '不计强旺与衰弱运用此诀可直断': '总论2.png',
    '此诀灵活运用好人称断命活神仙': '总论3.png',

    '甲木日干人刚强自律严谨好榜样': '甲1.png',
    '领导才能人称羡组织能力也很强': '甲2.png',
    '为人处事很正直多人信赖都夸奖': '甲3.png',
    '人逢甲木有疤痕不在脸上发中藏': '甲4.png',

    '乙木日干情脆弱外表内心不一样': '乙1.png',
    '外表懦弱又谨慎内心固执不改张': '乙2.png',
    '为人拘泥又保守常把他人细端详': '乙3.png',
    '人逢乙木筋骨疼易被棍棒碰损伤': '乙4.png',

    '丙火日干美敦厚文明有礼热心肠': '丙1.png',
    '活泼开朗乐天派人多之处善演讲': '丙2.png',
    '易被误解好表现缺乏毅力不久长': '丙3.png',
    '人逢丙日婚不顺痣痞定然在身上': '丙4.png',

    '丁火日干有礼貌言谈举止文明样': '丁1.png',
    '外柔掩盖上进心遇事缜密细思量': '丁2.png',
    '若遇挫折心不快一生想当常胜将': '丁3.png',
    '人逢丁火有伤残女人剖产取儿郎': '丁4.png',

    '戊土日干诚厚重善交朋友不久长': '戊1.png',
    '手腕灵活用心计与人好处彼全忘': '戊2.png',
    '人逢戊土喜打扮遇事常常无主张': '戊3.png',
    '胎印定在前腹部他柱临之细审详': '戊4.png',

    '己土日干人心细循规蹈矩小度量': '己1.png',
    '为人厚道能吃苦外表柔顺贪心强': '己2.png',
    '常常猜疑多苦恼同胞家人挂心上': '己3.png',
    '人逢己土有痣痞脾胃不佳欠营养': '己4.png',

    '庚金日干属纯阳为己享受愿奔忙': '庚1.png',
    '满腹文章口才好善于处事美名扬': '庚2.png',
    '勇敢好斗鸣不平坚强不屈粗野郎': '庚3.png',
    '易犯气管肺部病小心骨折易损伤': '庚4.png',

    '辛金日干人聪明性温质清人刚强': '辛1.png',
    '思想顽固持己见为人处事不张狂': '辛2.png',
    '善于动脑想办法排除困难事顺畅': '辛3.png',
    '筋骨疼痛口腔病时刻小心金属伤': '辛4.png',

    '壬水日干人聪明心口不一反无常': '壬1.png',
    '善于疏导凝聚力度量很大依赖强': '壬2.png',
    '性暴散乱无所谓常惹是非把财伤': '壬3.png',
    '易得肾病腰腿疼问医求诊记心上': '壬4.png',

    '癸水日干人勤勉正直廉洁楷模样': '癸1.png',
    '生命强旺能吃苦能屈能伸不争强': '癸2.png',
    '逆境之处拓新路忍辱负重人夸奖': '癸3.png',
    '肾脏肛门易得病女人妇科病应防': '癸4.png',
}


def compact(s: str) -> str:
    return re.sub(r'\s+', '', s).replace('　', '').replace(' ', '')


def insert_picture_after(doc, paragraph, img_path, width_cm=8.0):
    # Create at end, then move XML node after target paragraph.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(img_path, width=Cm(width_cm))
    paragraph._p.addnext(p._p)
    return p


def embed_word():
    doc = Document(IN_DOC)
    inserted = 0
    found = set()
    # iterate over a snapshot, because we insert new paragraphs as we go
    for p in list(doc.paragraphs):
        key = compact(p.text)
        if key in FORMULA_TO_IMAGE:
            img = FORMULA_TO_IMAGE[key]
            img_path = os.path.join(COMIC_DIR, img)
            if os.path.exists(img_path):
                insert_picture_after(doc, p, img_path, width_cm=8.2)
                inserted += 1
                found.add(key)
            else:
                print(f'[MISSING IMG] {img_path}')
    doc.save(OUT_DOC)
    missing = [v for k, v in FORMULA_TO_IMAGE.items() if k not in found]
    print(f'[WORD] inserted={inserted}, output={OUT_DOC}')
    if missing:
        print(f'[WORD] missing formula matches: {len(missing)}')
        for m in missing:
            print('  ', m)
    return inserted, missing


def patch_markdown():
    if not os.path.exists(MD_PATH):
        print(f'[MD] source not found: {MD_PATH}')
        return 0
    text = open(MD_PATH, 'r', encoding='utf-8').read()
    inserted = 0
    # Insert after quote/formula occurrences. Pandoc may use line breaks or spaces, so use direct formula fragments line by line.
    line_to_img = {
        '四柱排定看日干，日干断诀记心间。': '总论1.png',
        '不计强旺与衰弱，运用此诀可直断。': '总论2.png',
        '此诀灵活运用好，人称断命活神仙。': '总论3.png',
        '甲木日干人刚强，自律严谨好榜样。': '甲1.png',
        '领导才能人称羡，组织能力也很强。': '甲2.png',
        '为人处事很正直，多人信赖都夸奖。': '甲3.png',
        '人逢甲木有疤痕，不在脸上发中藏。': '甲4.png',
        '乙木日干情脆弱，外表内心不一样。': '乙1.png',
        '外表懦弱又谨慎，内心固执不改张。': '乙2.png',
        '为人拘泥又保守，常把他人细端详。': '乙3.png',
        '人逢乙木筋骨疼，易被棍棒碰损伤。': '乙4.png',
        '丙火日干美敦厚，文明有礼热心肠。': '丙1.png',
        '活泼开朗乐天派，人多之处善演讲。': '丙2.png',
        '易被误解好表现，缺乏毅力不久长。': '丙3.png',
        '人逢丙日婚不顺，痣痞定然在身上。': '丙4.png',
        '丁火日干有礼貌，言谈举止文明样。': '丁1.png',
        '外柔掩盖上进心，遇事缜密细思量。': '丁2.png',
        '若遇挫折心不快，一生想当常胜将。': '丁3.png',
        '人逢丁火有伤残，女人剖产取儿郎。': '丁4.png',
        '戊土日干诚厚重，善交朋友不久长。': '戊1.png',
        '手腕灵活用心计，与人好处彼全忘。': '戊2.png',
        '人逢戊土喜打扮，遇事常常无主张。': '戊3.png',
        '胎印定在前腹部，他柱临之细审详。': '戊4.png',
        '己土日干人心细，循规蹈矩小度量。': '己1.png',
        '为人厚道能吃苦，外表柔顺贪心强。': '己2.png',
        '常常猜疑多苦恼，同胞家人挂心上。': '己3.png',
        '人逢己土有痣痞，脾胃不佳欠营养。': '己4.png',
        '庚金日干属纯阳，为己享受愿奔忙。': '庚1.png',
        '满腹文章口才好，善于处事美名扬。': '庚2.png',
        '勇敢好斗鸣不平，坚强不屈粗野郎。': '庚3.png',
        '易犯气管肺部病，小心骨折易损伤。': '庚4.png',
        '辛金日干人聪明，性温质清人刚强。': '辛1.png',
        '思想顽固持己见，为人处事不张狂。': '辛2.png',
        '善于动脑想办法，排除困难事顺畅。': '辛3.png',
        '筋骨疼痛口腔病，时刻小心金属伤。': '辛4.png',
        '壬水日干人聪明，心口不一反无常。': '壬1.png',
        '善于疏导凝聚力，度量很大依赖强。': '壬2.png',
        '性暴散乱无所谓，常惹是非把财伤。': '壬3.png',
        '易得肾病腰腿疼，问医求诊记心上。': '壬4.png',
        '癸水日干人勤勉，正直廉洁楷模样。': '癸1.png',
        '生命强旺能吃苦，能屈能伸不争强。': '癸2.png',
        '逆境之处拓新路，忍辱负重人夸奖。': '癸3.png',
        '肾脏肛门易得病，女人妇科病应防。': '癸4.png',
    }
    for formula, img in line_to_img.items():
        if formula in text:
            rel = f'![[illustrations/comic/{img}]]'
            if rel not in text:
                text = text.replace(formula, formula + '\n\n' + rel, 1)
                inserted += 1
    open(OUT_MD, 'w', encoding='utf-8').write(text)
    print(f'[MD] inserted={inserted}, output={OUT_MD}')
    return inserted

if __name__ == '__main__':
    print(f'Comic image count: {len([f for f in os.listdir(COMIC_DIR) if f.lower().endswith(".png")])}')
    embed_word()
    patch_markdown()
